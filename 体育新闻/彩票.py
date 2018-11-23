import requests
from lxml import etree
import time
from docx import Document
import re

def info(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36 SE 2.X MetaSr 1.0'
    }
    res = requests.get(url, headers=headers)
    res.encoding = 'GBK'
    selector = etree.HTML(res.text)
    print(url)
    title = selector.xpath('//*[@id="container"]/div[1]/div[2]/h1/text()')
    comment = selector.xpath('//*[@id="contentText"]/div[1]/p/text()')
    comment = str(comment).replace('\\u3000','')
    comment = comment.replace('[\'', '')
    comment = comment.replace('\']', '')
    comment = comment.replace('\'', '')
    if len(comment) == 0:
        comment = selector.xpath('//*[@id="contentText"]/div[1]/div/text()')
        comment = str(comment).strip()
    #//*[@id="contentText"]/div[1]/div[1]/table/tbody/tr/td/img
    #//*[@id="contentText"]/div[1]/div/table/tbody/tr/td/img
    #//*[@id="contentText"]/div[1]/div[1]/text()[2]
    img = selector.xpath('//*[@id="contentText"]/div[1]/div/table/tbody/tr/td/img/@src')
    print(title, comment, img)
    try:
        if len(img) > 0:
            photo = requests.get(img[0], headers=headers)
            i = str(img[0]).split('/')[-1]
            fp = open('img/{}'.format(i), 'wb')
            fp.write(photo.content)
            fp.close()
            document = Document()
            document.add_paragraph(title[-1] + '\n' + comment)  # 向文档里添加文字
            document.add_picture('img/{}'.format(i))  # 向文档里添加图片
            document.save('新闻/{}.docx'.format(title[-1].replace(' ','')))  # 保存文档
        else:
            document = Document()
            document.add_paragraph(title[-1] + '\n' + comment)  # 向文档里添加文字
            document.save('新闻/{}.docx'.format(title[-1].replace(' ','')))  # 保存文档
    except:
        pass

def info_url(url):
    headers = {
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36 SE 2.X MetaSr 1.0'
    }
    res = requests.get(url, headers=headers)
    res.encoding = 'GBK'
    selector = etree.HTML(res.text)
    urls = selector.xpath('/html/body/div[1]/div[4]/div[1]/div[7]/ul/li/a/@href')
    for url in urls:
        info(url)

if __name__ == '__main__':
    urls = ['http://caipiao.sohu.com/lotto/tc/index_{}.shtml'.format(i)for i in range(211,100,-1)]
    for url in urls:
        info_url(url)
        time.sleep(3)